"""Word 文档解析器"""

import logging
from pathlib import Path
from typing import List

from docx import Document
from src.parsers.base import BaseParser, ParsedContent

logger = logging.getLogger(__name__)


class WordParser(BaseParser):
    """Word (.docx) 文档解析器"""

    def parse(self, file_path: Path) -> ParsedContent:
        """解析 Word 文档"""
        logger.info(f"[WordParser] 开始解析: {file_path}")

        try:
            doc = Document(file_path)

            # 提取标题（使用第一个段落，或文档属性中的标题）
            title = self._extract_title(doc)

            # 提取所有段落内容
            content_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    # 将段落样式转换为 Markdown 格式
                    content_parts.append(self._convert_paragraph(para))

            content = "\n\n".join(content_parts)

            # 提取图片
            images = self._extract_images(file_path)

            # 提取元数据
            metadata = self._extract_metadata(doc)

            result = ParsedContent(
                title=title,
                content=content,
                images=images,
                metadata=metadata,
            )

            logger.info(
                f"[WordParser] 解析完成 - 标题: {title}, 图片数: {len(images)}"
            )
            return result

        except Exception as e:
            logger.error(f"[WordParser] 解析失败: {e}")
            from src.exceptions import FileReadError
            raise FileReadError(str(file_path), str(e))

    def _extract_title(self, doc: Document) -> str:
        """提取标题"""
        # 尝试从文档属性获取标题
        if doc.core_properties.title:
            return doc.core_properties.title

        # 否则使用第一个非空段落作为标题
        for para in doc.paragraphs:
            if para.text.strip():
                text = para.text.strip()
                # 限制标题长度
                if len(text) > 50:
                    return text[:50] + "..."
                return text

        return "未命名文档"

    def _convert_paragraph(self, para) -> str:
        """将段落转换为 Markdown 格式"""
        text = para.text.strip()
        if not text:
            return ""

        # 获取段落样式
        style_name = para.style.name if para.style else ""

        # 处理标题
        if style_name.startswith("Heading"):
            level = style_name[-1]  # Heading1 -> 1
            return f"{'#' * int(level)} {text}"

        # 处理列表
        if style_name.startswith("List"):
            return f"- {text}"

        # 普通段落
        return text

    def _extract_images(self, file_path: Path) -> List[Path]:
        """提取图片引用（Word 解析器暂不支持图片提取）"""
        # Word 文档的图片提取较复杂，暂不支持
        return []

    def _extract_metadata(self, doc: Document) -> dict:
        """提取元数据"""
        props = doc.core_properties
        return {
            "author": props.author or "",
            "created": str(props.created) if props.created else "",
            "modified": str(props.modified) if props.modified else "",
            "comments": props.comments or "",
        }

    def supports(self, file_path: Path) -> bool:
        """判断是否支持该文件类型"""
        return file_path.suffix.lower() in [".docx", ".doc"]
